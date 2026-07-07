import os
from pypdf import PdfReader
from docx import Document
from utils.logger import logger

class ResumeParserError(Exception):
    """Custom exception class for resume parsing errors."""
    pass

def parse_pdf(file_path: str) -> str:
    """Parses text from a PDF file using pypdf."""
    logger.info(f"Parsing PDF file: {file_path}")
    try:
        reader = PdfReader(file_path)
        text = ""
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        
        text = text.strip()
        if not text:
            raise ResumeParserError("PDF file contains no readable text (it may be scanned or empty).")
        return text
    except ResumeParserError as e:
        raise e
    except Exception as e:
        logger.error(f"Error parsing PDF file {file_path}: {str(e)}")
        raise ResumeParserError(f"Failed to parse PDF file: {str(e)}")

def parse_docx(file_path: str) -> str:
    """Parses text from a DOCX file using python-docx."""
    logger.info(f"Parsing DOCX file: {file_path}")
    try:
        doc = Document(file_path)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text.append(cell.text)
                    
        full_text = "\n".join(text).strip()
        if not full_text:
            raise ResumeParserError("DOCX file contains no readable text.")
        return full_text
    except ResumeParserError as e:
        raise e
    except Exception as e:
        logger.error(f"Error parsing DOCX file {file_path}: {str(e)}")
        raise ResumeParserError(f"Failed to parse DOCX file: {str(e)}")

def parse_txt(file_path: str) -> str:
    """Parses text from a TXT file with encoding fallbacks."""
    logger.info(f"Parsing TXT file: {file_path}")
    encodings = ["utf-8", "latin-1", "utf-16", "cp1252"]
    
    for encoding in encodings:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                text = f.read().strip()
                if not text:
                    raise ResumeParserError("TXT file is empty.")
                return text
        except UnicodeDecodeError:
            continue
        except ResumeParserError as e:
            raise e
        except Exception as e:
            logger.error(f"Error reading TXT file {file_path} with encoding {encoding}: {str(e)}")
            raise ResumeParserError(f"Failed to parse TXT file: {str(e)}")
            
    raise ResumeParserError("Failed to decode TXT file: unsupported encoding.")

def parse_file(file_path: str) -> str:
    """
    Main entry point for file parsing.
    Determines the file type and routes to the appropriate parser.
    """
    if not os.path.exists(file_path):
        logger.error(f"File not found: {file_path}")
        raise FileNotFoundError(f"File not found: {file_path}")
        
    if os.path.getsize(file_path) == 0:
        logger.error(f"Empty file: {file_path}")
        raise ResumeParserError(f"File is empty: {file_path}")
        
    _, ext = os.path.splitext(file_path.lower())
    
    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        # Support docx; .doc is checked but only .docx is fully supported by python-docx
        if ext == ".doc":
            logger.warning(f"File format is .doc: {file_path}. python-docx only supports .docx.")
        return parse_docx(file_path)
    elif ext == ".txt":
        return parse_txt(file_path)
    else:
        logger.error(f"Unsupported file format: {ext} for file {file_path}")
        raise ResumeParserError(f"Unsupported file format '{ext}'. Supported formats: .pdf, .docx, .txt")
