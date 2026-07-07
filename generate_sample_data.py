import os
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.colors import HexColor
from docx import Document

def create_pdf_resume(path: str, name: str, details: dict):
    """Generates a clean, professional PDF resume using ReportLab."""
    doc = SimpleDocTemplate(path, pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'ResumeTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=22,
        textColor=HexColor("#1A365D"),
        spaceAfter=12
    )
    
    section_heading = ParagraphStyle(
        'SectionHeading',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=14,
        textColor=HexColor("#2B6CB0"),
        spaceBefore=10,
        spaceAfter=6,
        keepWithNext=True
    )
    
    body_style = ParagraphStyle(
        'ResumeBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=HexColor("#2D3748"),
        spaceAfter=8
    )
    
    bullet_style = ParagraphStyle(
        'ResumeBullet',
        parent=body_style,
        leftIndent=15,
        firstLineIndent=-10,
        spaceAfter=4
    )

    story = []
    
    # Header
    story.append(Paragraph(name, title_style))
    story.append(Paragraph(f"Email: {details['email']} | Phone: {details['phone']} | Address: {details['address']}", body_style))
    story.append(Spacer(1, 10))
    
    # Summary
    story.append(Paragraph("PROFESSIONAL SUMMARY", section_heading))
    story.append(Paragraph(details['summary'], body_style))
    story.append(Spacer(1, 8))
    
    # Skills
    story.append(Paragraph("SKILLS", section_heading))
    story.append(Paragraph(details['skills_text'], body_style))
    story.append(Spacer(1, 8))
    
    # Experience
    story.append(Paragraph("EXPERIENCE", section_heading))
    for exp in details['experience']:
        story.append(Paragraph(f"<b>{exp['role']}</b> - {exp['company']} ({exp['duration']})", body_style))
        for point in exp['points']:
            story.append(Paragraph(f"• {point}", bullet_style))
        story.append(Spacer(1, 4))
    story.append(Spacer(1, 8))
    
    # Education
    story.append(Paragraph("EDUCATION", section_heading))
    for edu in details['education']:
        story.append(Paragraph(f"<b>{edu['degree']}</b> - {edu['institution']} ({edu['year']})", body_style))
        if 'details' in edu:
            story.append(Paragraph(edu['details'], body_style))
            
    doc.build(story)

def create_docx_resume(path: str, name: str, details: dict):
    """Generates a clean DOCX resume using python-docx."""
    doc = Document()
    
    # Title / Name
    title = doc.add_heading(name, level=0)
    title.alignment = 0 # Left aligned
    
    # Contact Info
    contact = doc.add_paragraph()
    contact.add_run(f"Email: {details['email']} | Phone: {details['phone']} | Address: {details['address']}")
    
    # Summary
    doc.add_heading("PROFESSIONAL SUMMARY", level=1)
    doc.add_paragraph(details['summary'])
    
    # Skills
    doc.add_heading("SKILLS", level=1)
    doc.add_paragraph(details['skills_text'])
    
    # Experience
    doc.add_heading("EXPERIENCE", level=1)
    for exp in details['experience']:
        p = doc.add_paragraph()
        p.add_run(f"{exp['role']} at {exp['company']} ({exp['duration']})").bold = True
        for point in exp['points']:
            doc.add_paragraph(point, style='List Bullet')
            
    # Education
    doc.add_heading("EDUCATION", level=1)
    for edu in details['education']:
        p = doc.add_paragraph()
        p.add_run(f"{edu['degree']} from {edu['institution']} ({edu['year']})").bold = True
        if 'details' in edu:
            doc.add_paragraph(edu['details'])
            
    doc.save(path)

def create_txt_resume(path: str, name: str, details: dict):
    """Generates a structured plain text resume."""
    with open(path, "w", encoding="utf-8") as f:
        f.write(f"{name}\n")
        f.write(f"Email: {details['email']} | Phone: {details['phone']} | Address: {details['address']}\n")
        f.write("="*50 + "\n\n")
        
        f.write("PROFESSIONAL SUMMARY\n")
        f.write("-"*20 + "\n")
        f.write(f"{details['summary']}\n\n")
        
        f.write("SKILLS\n")
        f.write("-"*20 + "\n")
        f.write(f"{details['skills_text']}\n\n")
        
        f.write("EXPERIENCE\n")
        f.write("-"*20 + "\n")
        for exp in details['experience']:
            f.write(f"{exp['role']} - {exp['company']} ({exp['duration']})\n")
            for point in exp['points']:
                f.write(f"- {point}\n")
            f.write("\n")
            
        f.write("EDUCATION\n")
        f.write("-"*20 + "\n")
        for edu in details['education']:
            f.write(f"{edu['degree']} - {edu['institution']} ({edu['year']})\n")
            if 'details' in edu:
                f.write(f"{edu['details']}\n")
            f.write("\n")

def main():
    print("Generating sample job description and resumes...")
    os.makedirs("jd", exist_ok=True)
    os.makedirs("resumes", exist_ok=True)
    os.makedirs("output", exist_ok=True)
    
    # 1. Write the Job Description
    jd_content = """JUNIOR AI RESEARCH ASSOCIATE
Location: Remote / Hybrid (New York, NY)
Type: Full-Time

POSITION OVERVIEW:
We are seeking a highly motivated and technically proficient Junior AI Research Associate to join our advanced AI R&D team. In this role, you will assist in developing, optimizing, and evaluating state-of-the-art NLP systems, large language models (LLMs), and semantic search engines. The ideal candidate has strong coding capabilities in Python, hands-on experience with deep learning frameworks, and a solid understanding of natural language processing concepts.

KEY RESPONSIBILITIES:
- Implement, train, and fine-tune NLP models, specifically Sentence Transformers, BERT-based architectures, and LLMs.
- Build and evaluate semantic similarity, retrieval-augmented generation (RAG), and text classification systems.
- Conduct data preprocessing, text cleaning, parsing, and analysis of large, unstructured text corpora.
- Collaborate with senior scientists to write research prototypes, technical reports, and documentation.
- Maintain and deploy lightweight demo applications using frameworks like Streamlit.

REQUIRED SKILLS & QUALIFICATIONS:
- Python programming expertise (object-oriented, clean code principles, PEP 8).
- Experience with PyTorch or TensorFlow, and machine learning libraries (Scikit-Learn, Pandas, NumPy).
- Direct experience with Natural Language Processing (NLP) concepts (embeddings, cosine similarity, tokenization).
- Familiarity with Hugging Face transformers library and Sentence Transformers (e.g., all-MiniLM-L6-v2).
- Strong research capabilities: ability to read scientific papers, extract methodologies, and implement prototypes.
- Excellent communication and technical writing skills.
- Bachelor's or Master's degree in Computer Science, Data Science, Computational Linguistics, or a related quantitative field.
"""
    with open("jd/job_description.txt", "w", encoding="utf-8") as f:
        f.write(jd_content)
    print("Created jd/job_description.txt")

    # 2. Resumes database
    candidates = [
        # Candidate 1: Alice Smith (Strong Match, PDF)
        {
            "name": "Alice Smith",
            "format": "pdf",
            "details": {
                "email": "alice.smith@email.com",
                "phone": "555-0192",
                "address": "Boston, MA",
                "summary": "Enthusiastic Junior ML Engineer specializing in Natural Language Processing and Deep Learning. Experienced in implementing Sentence Transformers and BERT-based semantic search tools. Proven record of conducting text mining research and writing clean, modular Python prototypes.",
                "skills_text": "Technical Skills: Python (expert), PyTorch, Hugging Face Transformers, Sentence Transformers (all-MiniLM-L6-v2, MPNet), Pandas, NumPy, Scikit-Learn, Git, SQL, Streamlit, NLP, Semantic Search, BERT, Text Preprocessing, RAG.",
                "experience": [
                    {
                        "role": "NLP Research Intern",
                        "company": "DeepText Labs",
                        "duration": "June 2025 - Present",
                        "points": [
                            "Developed and deployed an internal semantic search tool using Sentence Transformers (all-MiniLM-L6-v2) and Pinecone vector database, improving retrieval accuracy by 25%.",
                            "Conducted preprocessing, embedding extraction, and text similarity evaluation on over 100k scientific documents.",
                            "Created interactive demo dashboards using Streamlit to showcase RAG system capabilities to stakeholders.",
                            "Implemented object-oriented Python modules for text tokenization, filtering, and model evaluation."
                        ]
                    }
                ],
                "education": [
                    {
                        "degree": "M.S. in Computer Science (Concentration in AI)",
                        "institution": "Boston University",
                        "year": "2025",
                        "details": "Relevant Coursework: Advanced Natural Language Processing, Machine Learning, Deep Learning, Statistics. Thesis on semantic embedding fine-tuning."
                    }
                ]
            }
        },
        # Candidate 2: Bob Johnson (Strong Match, DOCX)
        {
            "name": "Bob Johnson",
            "format": "docx",
            "details": {
                "email": "bob.johnson@email.com",
                "phone": "555-0143",
                "address": "Seattle, WA",
                "summary": "Recent Data Science graduate with hands-on experience in machine learning and NLP. Highly proficient in Python, PyTorch, and deploying custom models. Eager to contribute to bleeding-edge AI research and product development.",
                "skills_text": "Skills: Python, PyTorch, Git, Pandas, NumPy, Scikit-Learn, Hugging Face, Transformers, NLP, Tokenization, Cosine Similarity, Vector Databases, FastAPI, Docker, SQL, LaTeX.",
                "experience": [
                    {
                        "role": "AI Research Assistant",
                        "company": "University AI Research Lab",
                        "duration": "Sep 2024 - May 2025",
                        "points": [
                            "Implemented and compared multiple sentence embedding models (including all-MiniLM-L6-v2) for question-answering systems.",
                            "Preprocessed large unstructured text datasets from medical literature using regular expressions, NLTK, and SpaCy.",
                            "Co-authored two technical reports summarizing model benchmark results and performance trade-offs."
                        ]
                    }
                ],
                "education": [
                    {
                        "degree": "B.S. in Data Science",
                        "institution": "University of Washington",
                        "year": "2024",
                        "details": "Graduated with Honors. Focus on Machine Learning and Computational Linguistics."
                    }
                ]
            }
        },
        # Candidate 3: Charlie Brown (Moderate Match, TXT)
        {
            "name": "Charlie Brown",
            "format": "txt",
            "details": {
                "email": "charlie.brown@email.com",
                "phone": "555-0162",
                "address": "Chicago, IL",
                "summary": "Data Analyst with 2 years of professional experience writing Python scripts, querying databases, and building statistical models. Transitioning into machine learning with a focus on NLP and intelligent text processing.",
                "skills_text": "Technical Skills: Python, SQL, Pandas, NumPy, Scikit-Learn, Matplotlib, Tableau, Git, Basic Machine Learning, Statistics, Regression, Clustering.",
                "experience": [
                    {
                        "role": "Junior Data Analyst",
                        "company": "RetailMetrics Inc.",
                        "duration": "July 2024 - Present",
                        "points": [
                            "Wrote reusable Python scripts to automate daily ETL data pipelines, processing millions of customer feedback entries.",
                            "Performed sentiment analysis on customer reviews using Scikit-Learn and VADER, providing actionable insights to product managers.",
                            "Maintained clean documentation and version control using Git across all analytic scripts."
                        ]
                    }
                ],
                "education": [
                    {
                        "degree": "B.S. in Statistics",
                        "institution": "University of Illinois",
                        "year": "2023"
                    }
                ]
            }
        },
        # Candidate 4: Diana Prince (Moderate Match, PDF)
        {
            "name": "Diana Prince",
            "format": "pdf",
            "details": {
                "email": "diana.prince@email.com",
                "phone": "555-0199",
                "address": "Washington, DC",
                "summary": "Research Specialist with a background in Computational Linguistics and cognitive systems. Experienced in academic research, analyzing text corpora, and prototyping linguistic models in Python.",
                "skills_text": "Skills: Python, NLTK, SpaCy, Corpus Linguistics, Research Writing, Pandas, NumPy, Git, Statistics, Qualitative Text Analysis, LaTeX, Technical Documentation.",
                "experience": [
                    {
                        "role": "Research Associate (Linguistics)",
                        "company": "Cognitive Science Institute",
                        "duration": "Oct 2024 - Present",
                        "points": [
                            "Curated and annotated semantic text corpora for syntax and semantics research, using NLTK and custom regex parser scripts.",
                            "Conducted comprehensive literature reviews on transformer model evaluations and summarized findings for the core engineering team.",
                            "Drafted weekly research summaries and contributed to drafting papers for ACL conferences."
                        ]
                    }
                ],
                "education": [
                    {
                        "degree": "M.A. in Computational Linguistics",
                        "institution": "Georgetown University",
                        "year": "2024"
                    }
                ]
            }
        },
        # Candidate 5: Evan Wright (Weak Match, DOCX)
        {
            "name": "Evan Wright",
            "format": "docx",
            "details": {
                "email": "evan.wright@email.com",
                "phone": "555-0187",
                "address": "San Francisco, CA",
                "summary": "Dynamic Frontend Developer with 3 years of experience crafting interactive web applications. Expertise in React, JavaScript, and CSS. Looking to expand frontend development skills into AI-related dashboards.",
                "skills_text": "Skills: JavaScript, TypeScript, React, HTML5, CSS3, TailwindCSS, Node.js, Express, Webpack, Git, REST APIs, UI/UX Design.",
                "experience": [
                    {
                        "role": "Frontend Web Developer",
                        "company": "WebFlow Solutions",
                        "duration": "Jan 2023 - Present",
                        "points": [
                            "Built and optimized 20+ responsive web pages using React and TailwindCSS, improving load times by 40%.",
                            "Integrated RESTful APIs to display real-time metrics and charts in interactive developer portals.",
                            "Collaborated with UI/UX designers to implement clean, modern components and dark mode interfaces."
                        ]
                    }
                ],
                "education": [
                    {
                        "degree": "B.S. in Web Design and Development",
                        "institution": "San Jose State University",
                        "year": "2022"
                    }
                ]
            }
        },
        # Candidate 6: Fiona Gallagher (Strong Match, TXT)
        {
            "name": "Fiona Gallagher",
            "format": "txt",
            "details": {
                "email": "fiona.g@email.com",
                "phone": "555-0112",
                "address": "Austin, TX",
                "summary": "Junior AI Developer specializing in PyTorch-driven NLP applications. Strong knowledge of Hugging Face pipelines, transformer fine-tuning, and semantic similarity evaluations. Passionate about clean code and rapid prototyping.",
                "skills_text": "Keywords: Python, PyTorch, Transformers, Hugging Face, Sentence Transformers (all-MiniLM-L6-v2), BERT, Tokenization, Cosine Similarity, Git, GitHub, Pandas, Streamlit, NLP, FastAPI, Linux.",
                "experience": [
                    {
                        "role": "AI Engineer (Contractor)",
                        "company": "NeuralCraft Labs",
                        "duration": "Nov 2024 - May 2025",
                        "points": [
                            "Implemented a sentence-similarity search engine utilizing Hugging Face Sentence Transformers and PyTorch, processing query similarity matches under 50ms.",
                            "Assisted in training and evaluating custom tokenizer models on domain-specific vocabulary.",
                            "Built a Streamlit dashboard showing model performance, classification metrics, and attention weight visualizations."
                        ]
                    }
                ],
                "education": [
                    {
                        "degree": "B.S. in Computer Science",
                        "institution": "University of Texas at Austin",
                        "year": "2024",
                        "details": "Completed capstone project on AI-driven text summarization using transformer networks."
                    }
                ]
            }
        },
        # Candidate 7: George Costanza (Weak Match, PDF)
        {
            "name": "George Costanza",
            "format": "pdf",
            "details": {
                "email": "george.c@email.com",
                "phone": "555-0133",
                "address": "Queens, NY",
                "summary": "Experienced IT Project Manager and Scrum Master. Skilled in Agile methodologies, coordinating cross-functional technical teams, managing sprints, and communicating status reports to executives.",
                "skills_text": "Skills: Agile, Scrum, Jira, Confluence, Project Management, SDLC, Risk Management, Communication, Team Leadership, MS Office.",
                "experience": [
                    {
                        "role": "Technical Project Manager",
                        "company": "Vandelay Industries",
                        "duration": "Mar 2023 - Present",
                        "points": [
                            "Led daily stand-ups, sprint planning, and retrospective meetings for a team of 8 software engineers.",
                            "Optimized delivery pipeline using Jira workflows, cutting project delivery times by 15%.",
                            "Drafted weekly executive summaries outlining project timeline, risks, and resource allocation."
                        ]
                    }
                ],
                "education": [
                    {
                        "degree": "B.A. in Business Administration",
                        "institution": "Queens College",
                        "year": "2018"
                    }
                ]
            }
        },
        # Candidate 8: Hannah Abbott (Moderate Match, DOCX)
        {
            "name": "Hannah Abbott",
            "format": "docx",
            "details": {
                "email": "hannah.a@email.com",
                "phone": "555-0155",
                "address": "Denver, CO",
                "summary": "Python Developer with 2 years of backend programming experience. Expert in Flask, FastAPI, and database operations. Eager to transition into AI research and apply software engineering best practices to model deployment.",
                "skills_text": "Technical Skills: Python (expert), Flask, FastAPI, SQL, PostgreSQL, Git, Docker, Unit Testing, NumPy, Pandas, Scikit-Learn, REST APIs, OOP, PEP 8.",
                "experience": [
                    {
                        "role": "Junior Backend Developer",
                        "company": "RockyMountain Software",
                        "duration": "June 2024 - Present",
                        "points": [
                            "Developed and maintained clean, test-driven REST APIs using FastAPI and SQLAlchemy, handling 50k requests daily.",
                            "Dockerized backend microservices to ensure consistent environment setups across local development and staging servers.",
                            "Refactored legacy Python scripts to align with PEP 8 standards, introducing static type hinting and comprehensive docstrings."
                        ]
                    }
                ],
                "education": [
                    {
                        "degree": "B.S. in Software Engineering",
                        "institution": "Colorado State University",
                        "year": "2024"
                    }
                ]
            }
        },
        # Candidate 9: Ian Malcolm (Weak Match, TXT)
        {
            "name": "Ian Malcolm",
            "format": "txt",
            "details": {
                "email": "chaos.ian@email.com",
                "phone": "555-0167",
                "address": "Santa Fe, NM",
                "summary": "Quantitative Analyst and Theoretical Statistician. Expert in modeling complex non-linear systems, chaos theory, statistical simulations, and mathematical modeling.",
                "skills_text": "Skills: R, MATLAB, Mathematica, LaTeX, Statistical Modeling, Probability Theory, Data Visualization, Complex Systems Analysis.",
                "experience": [
                    {
                        "role": "Senior Statistical Consultant",
                        "company": "InGen Dynamics",
                        "duration": "Aug 2021 - Present",
                        "points": [
                            "Modeled population dynamics and ecological trends using complex statistical frameworks in R and MATLAB.",
                            "Wrote extensive academic research papers detailing model findings and presented them at international math conferences.",
                            "Consulted engineering teams on experimental design and risk assessment models."
                        ]
                    }
                ],
                "education": [
                    {
                        "degree": "Ph.D. in Applied Mathematics",
                        "institution": "University of Texas",
                        "year": "2020"
                    }
                ]
            }
        },
        # Candidate 10: Julia Roberts (Strong Match, PDF)
        {
            "name": "Julia Roberts",
            "format": "pdf",
            "details": {
                "email": "julia.r@email.com",
                "phone": "555-0101",
                "address": "Atlanta, GA",
                "summary": "Ph.D. Graduate in Computational Linguistics with extensive research in Natural Language Processing, Transformer architectures, and Semantic Retrieval. Proven track record of publications in top AI venues (ACL, EMNLP) and deep hands-on expertise in Python and PyTorch.",
                "skills_text": "Skills: Python, PyTorch, NLP, Transformers, Sentence Transformers, Cosine Similarity, BERT, LLMs, Research Prototypes, Academic Writing, Pandas, NumPy, Scikit-Learn, Git, LaTeX.",
                "experience": [
                    {
                        "role": "AI Research Fellow",
                        "company": "Language Research Labs",
                        "duration": "Sep 2023 - May 2026",
                        "points": [
                            "Researched and implemented state-of-the-art embedding alignment techniques for multilingual Sentence Transformers, boosting cross-lingual semantic search retrieval by 18%.",
                            "Authored 3 publications in peer-reviewed NLP conferences demonstrating novel token-pruning methodologies.",
                            "Developed custom deep learning pipelines using PyTorch and Hugging Face to evaluate semantic textual similarity."
                        ]
                    }
                ],
                "education": [
                    {
                        "degree": "Ph.D. in Computational Linguistics",
                        "institution": "Georgia Institute of Technology",
                        "year": "2026",
                        "details": "Dissertation on 'Fine-Grained Semantic Representation in Dense Vector Spaces'. Received outstanding research award."
                    }
                ]
            }
        }
    ]
    
    # Generate candidates in their respective formats
    for i, cand in enumerate(candidates, 1):
        name = cand["name"]
        fmt = cand["format"]
        details = cand["details"]
        
        # Format filename
        filename = f"resume_{i:02d}_{name.lower().replace(' ', '_')}.{fmt}"
        filepath = os.path.join("resumes", filename)
        
        if fmt == "pdf":
            create_pdf_resume(filepath, name, details)
        elif fmt == "docx":
            create_docx_resume(filepath, name, details)
        elif fmt == "txt":
            create_txt_resume(filepath, name, details)
            
        print(f"Generated {filepath}")
        
    print("Successfully generated all sample data!")

if __name__ == "__main__":
    main()
