import os
import pandas as pd
import streamlit as st
import numpy as np
import json
from typing import List, Dict, Any, Tuple
import io

# Import local utilities
from utils.logger import logger
from utils.parser import parse_file, ResumeParserError
from utils.extractor import extract_resume_info
from utils.scorer import ResumeScorer

# -------------------------------------------------------------
# Configuration & Theme Setup
# -------------------------------------------------------------
st.set_page_config(
    page_title="AI Resume Screening Agent",
    page_icon="💼",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom Premium CSS for modern glassmorphism UI
st.markdown("""
<style>
    /* Main container styling */
    .main {
        background-color: #f7fafc;
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
    }
    
    /* Dark mode adjustments */
    @media (prefers-color-scheme: dark) {
        .main {
            background-color: #0f172a;
        }
    }
    
    /* Header card */
    .header-card {
        background: linear-gradient(135deg, #1e3a8a 0%, #3b82f6 100%);
        color: white;
        padding: 2.5rem;
        border-radius: 16px;
        box-shadow: 0 10px 15px -3px rgba(0, 0, 0, 0.1), 0 4px 6px -2px rgba(0, 0, 0, 0.05);
        margin-bottom: 2rem;
    }
    .header-card h1 {
        color: white !important;
        margin: 0;
        font-size: 2.5rem;
        font-weight: 800;
    }
    .header-card p {
        margin-top: 0.5rem;
        opacity: 0.9;
        font-size: 1.1rem;
    }
    
    /* Metric container style */
    .metric-card {
        background: white;
        padding: 1.5rem;
        border-radius: 12px;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05), 0 2px 4px -1px rgba(0, 0, 0, 0.03);
        border: 1px solid #e2e8f0;
        text-align: center;
    }
    @media (prefers-color-scheme: dark) {
        .metric-card {
            background: #1e293b;
            border-color: #334155;
            color: white;
        }
    }
    
    /* Highlighting matching text */
    .strength-highlight {
        background-color: #def7ec;
        color: #03543f;
        padding: 2px 6px;
        border-radius: 4px;
        font-weight: 500;
    }
    
    /* Custom button formatting */
    div.stButton > button:first-child {
        background: linear-gradient(135deg, #2563eb 0%, #1d4ed8 100%);
        color: white;
        border: none;
        padding: 0.75rem 2rem;
        border-radius: 8px;
        font-weight: 600;
        transition: all 0.2s ease-in-out;
        box-shadow: 0 4px 6px -1px rgba(37, 99, 235, 0.2);
    }
    div.stButton > button:first-child:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 12px -1px rgba(37, 99, 235, 0.3);
    }
</style>
""", unsafe_allow_html=True)

# -------------------------------------------------------------
# Lazy-loaded Cached Scorer
# -------------------------------------------------------------
@st.cache_resource
def get_scorer() -> ResumeScorer:
    """Instantiates and caches the SentenceTransformer ResumeScorer."""
    return ResumeScorer()

# -------------------------------------------------------------
# File Processing Helpers
# -------------------------------------------------------------
def parse_uploaded_file(uploaded_file) -> str:
    """Parses text from streamlit file_uploader file in-memory."""
    filename = uploaded_file.name
    logger.info(f"Parsing in-memory uploaded file: {filename}")
    
    if uploaded_file.size == 0:
        raise ResumeParserError(f"Uploaded file '{filename}' is empty.")
        
    _, ext = os.path.splitext(filename.lower())
    
    # Read bytes
    file_bytes = uploaded_file.read()
    
    if ext == ".pdf":
        from pypdf import PdfReader
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            text = text.strip()
            if not text:
                raise ResumeParserError("PDF contains no readable text (likely scanned image).")
            return text
        except Exception as e:
            raise ResumeParserError(f"PDF parsing failed: {e}")
            
    elif ext in [".docx", ".doc"]:
        from docx import Document
        try:
            doc = Document(io.BytesIO(file_bytes))
            text = []
            for p in doc.paragraphs:
                text.append(p.text)
            for t in doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        text.append(cell.text)
            full_text = "\n".join(text).strip()
            if not full_text:
                raise ResumeParserError("DOCX contains no readable text.")
            return full_text
        except Exception as e:
            raise ResumeParserError(f"DOCX parsing failed: {e}")
            
    elif ext == ".txt":
        # Try various decoders
        for encoding in ["utf-8", "latin-1", "utf-16", "cp1252"]:
            try:
                return file_bytes.decode(encoding).strip()
            except UnicodeDecodeError:
                continue
        raise ResumeParserError("Failed to decode TXT file: unsupported encoding.")
        
    else:
        raise ResumeParserError(f"Unsupported file format: {ext}")

# -------------------------------------------------------------
# Main Application Flow
# -------------------------------------------------------------
def main():
    # Render Hero Header
    st.markdown("""
    <div class="header-card">
        <h1>AI Resume Screening Agent</h1>
        <p>Rank and analyze candidate resumes against job descriptions with semantic NLP intelligence</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Initialize variables in session state
    if "results_df" not in st.session_state:
        st.session_state.results_df = None
    if "detailed_results" not in st.session_state:
        st.session_state.detailed_results = {}
    if "job_description_text" not in st.session_state:
        st.session_state.job_description_text = ""
        
    # Sidebar control panel
    st.sidebar.header("🔧 Configuration Control")
    
    jd_input_method = st.sidebar.radio("Job Description Input", ["Paste Text", "Upload JD File"])
        
    st.sidebar.divider()
    
    # -------------------------------------------------------------
    # Get Job Description Text
    # -------------------------------------------------------------
    jd_text = ""
    # Custom input
    if jd_input_method == "Paste Text":
        jd_text = st.sidebar.text_area("Paste Job Description here:", height=300, value=st.session_state.job_description_text)
        st.session_state.job_description_text = jd_text
    else:
        uploaded_jd = st.sidebar.file_uploader("Upload Job Description (.txt)", type=["txt"])
        if uploaded_jd:
            jd_text = uploaded_jd.read().decode("utf-8")
            st.sidebar.success(f"Uploaded: {uploaded_jd.name}")
            st.session_state.job_description_text = jd_text
                
    st.sidebar.divider()
    
    # -------------------------------------------------------------
    # Get Resumes
    # -------------------------------------------------------------
    uploaded_resumes = st.sidebar.file_uploader(
        "Upload Candidate Resumes (.pdf, .docx, .txt)",
        type=["pdf", "docx", "txt"],
        accept_multiple_files=True
    )
    if uploaded_resumes:
        st.sidebar.info(f"{len(uploaded_resumes)} resumes uploaded.")
            
    st.sidebar.divider()
    
    # Screening Action Button
    run_screening = st.sidebar.button("🚀 Run Screening Agent", use_container_width=True)
    
    # -------------------------------------------------------------
    # Main View Panels
    # -------------------------------------------------------------
    col_jd, col_dashboard = st.columns([1, 2])
    
    with col_jd:
        st.subheader("📋 Target Job Description")
        if jd_text:
            st.info("Loaded Job Description successfully. Showing excerpt below:")
            st.text_area("Job Description Preview", jd_text, height=450, disabled=True)
        else:
            st.warning("Please configure or paste a Job Description in the sidebar.")
            
    with col_dashboard:
        st.subheader("📊 Screening Status & Dashboard")
        
        if run_screening:
            if not jd_text.strip():
                st.error("Cannot run screening: Job Description is empty or missing.")
                return
                
            if not uploaded_resumes:
                st.error("Cannot run screening: No resumes uploaded.")
                return
                
            # Run Process
            logger.info("Initializing resume screening execution cycle...")
            scorer = get_scorer()
            
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            detailed_results = {}
            flat_results = []
            
            parsed_count = 0
            failed_count = 0
            
            for i, resume in enumerate(uploaded_resumes):
                # Update progress
                progress = (i) / len(uploaded_resumes)
                progress_bar.progress(progress)
                
                # Uploaded files
                file_name = resume.name
                status_text.text(f"Processing ({i+1}/{len(uploaded_resumes)}): {file_name}")
                try:
                    raw_text = parse_uploaded_file(resume)
                    # We pass file_name as dummy path
                    info = extract_resume_info(file_name, raw_text)
                    score_info = scorer.score_candidate(jd_text, info)
                    
                    detailed_results[info["candidate_name"]] = {
                        "info": info,
                        "scores": score_info
                    }
                    
                    flat_results.append({
                        "Candidate Name": info["candidate_name"],
                        "File Name": file_name,
                        "Score": score_info["overall_score"],
                        "Skills": ", ".join(info["skills_list"]) if info["skills_list"] else "Not Specified",
                        "Experience": info["experience_text"][:200] + ("..." if len(info["experience_text"]) > 200 else ""),
                        "Education": info["education_text"][:200] + ("..." if len(info["education_text"]) > 200 else ""),
                        "Reason for ranking": score_info["reasoning"]
                    })
                    parsed_count += 1
                except Exception as e:
                    logger.error(f"Failed to process uploaded file {file_name}: {e}")
                    failed_count += 1
                    flat_results.append({
                        "Candidate Name": f"Corrupted Candidate ({file_name})",
                        "File Name": file_name,
                        "Score": 0.0,
                        "Skills": "N/A",
                        "Experience": "N/A",
                        "Education": "N/A",
                        "Reason for ranking": f"PARSE ERROR: {str(e)}"
                    })
            
            progress_bar.progress(1.0)
            status_text.text("Screening cycle completed successfully!")
            
            # Form DataFrame
            df = pd.DataFrame(flat_results)
            df = df.sort_values(by="Score", ascending=False).reset_index(drop=True)
            df.index = df.index + 1
            df.index.name = "Rank"
            
            # Save variables in session state
            st.session_state.results_df = df.reset_index()
            st.session_state.detailed_results = detailed_results
            
            # Export to output/ folders automatically
            os.makedirs("output", exist_ok=True)
            csv_path = "output/results.csv"
            json_path = "output/results.json"
            
            df.to_csv(csv_path)
            
            # Create json export
            json_data = {
                "summary": {
                    "total_candidates": len(flat_results),
                    "parsed_successfully": parsed_count,
                    "failed_count": failed_count,
                    "average_score": round(df["Score"].mean(), 2) if len(df) > 0 else 0
                },
                "candidates": df.to_dict(orient="records")
            }
            with open(json_path, "w", encoding="utf-8") as f:
                json.dump(json_data, f, indent=4, default=str)
                
            st.success(f"Results successfully exported to: {csv_path} and {json_path}")
            
        # Display Dashboard Metrics & Results Table
        if st.session_state.results_df is not None:
            df = st.session_state.results_df
            total_cand = len(df)
            avg_score = round(df["Score"].mean(), 2)
            
            top_candidate = df.iloc[0]["Candidate Name"] if total_cand > 0 else "N/A"
            top_score = df.iloc[0]["Score"] if total_cand > 0 else 0
            
            # Summary metrics
            m_col1, m_col2, m_col3 = st.columns(3)
            with m_col1:
                st.markdown(f"""
                <div class="metric-card">
                    <p style="font-size:0.9rem; color:#718096; margin:0;">Total Screened</p>
                    <h2 style="font-size:2.2rem; font-weight:800; margin:5px 0 0 0; color:#2b6cb0;">{total_cand}</h2>
                </div>
                """, unsafe_allow_html=True)
            with m_col2:
                st.markdown(f"""
                <div class="metric-card">
                    <p style="font-size:0.9rem; color:#718096; margin:0;">Average Score</p>
                    <h2 style="font-size:2.2rem; font-weight:800; margin:5px 0 0 0; color:#2f855a;">{avg_score}%</h2>
                </div>
                """, unsafe_allow_html=True)
            with m_col3:
                st.markdown(f"""
                <div class="metric-card">
                    <p style="font-size:0.9rem; color:#718096; margin:0;">Top Candidate</p>
                    <h2 style="font-size:1.6rem; font-weight:800; margin:13px 0 0 0; color:#b7791f; text-overflow:ellipsis; overflow:hidden; white-space:nowrap;">{top_candidate} ({top_score}%)</h2>
                </div>
                """, unsafe_allow_html=True)
                
            st.divider()
            
            # Chart Visualization
            st.subheader("📈 Candidate Score Distribution")
            chart_df = df[["Candidate Name", "Score"]].copy()
            chart_df = chart_df.set_index("Candidate Name")
            st.bar_chart(chart_df, color="#3b82f6")
            
            st.divider()
            
            # Rank table
            st.subheader("🏆 Ranked Candidate Results")
            st.dataframe(
                df.set_index("Rank")[["Candidate Name", "Score", "Skills", "Reason for ranking"]],
                use_container_width=True
            )
            
            # Download section
            st.subheader("📥 Download Screening Reports")
            dl_col1, dl_col2 = st.columns(2)
            
            # Prep download data
            csv_buffer = io.StringIO()
            df.to_csv(csv_buffer, index=False)
            csv_data = csv_buffer.getvalue()
            
            json_data = json.dumps({
                "candidates": df.to_dict(orient="records")
            }, indent=4)
            
            with dl_col1:
                st.download_button(
                    label="Download Results CSV File",
                    data=csv_data,
                    file_name="results.csv",
                    mime="text/csv",
                    use_container_width=True
                )
            with dl_col2:
                st.download_button(
                    label="Download Results JSON File",
                    data=json_data,
                    file_name="results.json",
                    mime="application/json",
                    use_container_width=True
                )
        else:
            st.info("Upload resumes and a job description, then click 'Run Screening Agent' in the sidebar to start.")

    # -------------------------------------------------------------
    # Candidate Deep Dive Explorer
    # -------------------------------------------------------------
    if st.session_state.results_df is not None and st.session_state.detailed_results:
        st.divider()
        st.header("🔍 Candidate Deep Dive Explorer")
        
        candidates_list = list(st.session_state.detailed_results.keys())
        selected_candidate = st.selectbox("Select a Candidate to analyze details", candidates_list)
        
        if selected_candidate:
            cand_data = st.session_state.detailed_results[selected_candidate]
            info = cand_data["info"]
            scores = cand_data["scores"]
            
            exp_col1, exp_col2 = st.columns([1, 1])
            
            with exp_col1:
                st.markdown(f"### candidate: **{selected_candidate}**")
                st.write(f"**Source File:** {info['file_name']}")
                
                # Show parsed structured segments
                st.subheader("🎯 Parsed Key Sections")
                with st.expander("Skills Section Content", expanded=True):
                    st.write(info["skills_text"])
                with st.expander("Work Experience Content"):
                    st.write(info["experience_text"])
                with st.expander("Education Details"):
                    st.write(info["education_text"])
                    
            with exp_col2:
                st.subheader("⚖️ NLP Match Breakdown")
                
                # Metric display
                sc_col1, sc_col2, sc_col3 = st.columns(3)
                sc_col1.metric("Overall Score", f"{scores['overall_score']}%")
                sc_col2.metric("Skills Semantic Alignment", f"{round(scores['skills_similarity']*100, 1)}%")
                sc_col3.metric("Experience Alignment", f"{round(scores['experience_similarity']*100, 1)}%")
                
                st.write("**Overall Semantic Cosine Similarity (JD vs. Full Text):**", scores["overall_similarity"])
                
                st.subheader("💡 Screening Recommendation & Rationale")
                st.info(scores["reasoning"])
                
                # Show all scanned skills keywords
                st.subheader("🔑 Scanned Keyword Highlights")
                if info["skills_list"]:
                    # Create html badges
                    badges_html = " ".join([
                        f'<span class="strength-highlight" style="margin-right:8px; display:inline-block; margin-bottom:8px;">{skill}</span>' 
                        for skill in info["skills_list"]
                    ])
                    st.markdown(badges_html, unsafe_allow_html=True)
                else:
                    st.write("No common tech stack keywords identified.")
                    
                with st.expander("Full Text View"):
                    st.text(info["full_text"])

if __name__ == "__main__":
    main()
